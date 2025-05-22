import os
import csv
import json
import subprocess
from typing import List, Dict, Tuple, Optional
from pathlib import Path
from tqdm import tqdm
import pandas as pd
from docx import Document
import re
import yaml
from PyPDF2 import PdfReader
import xlrd
from datetime import datetime
import requests
import time
import unicodedata
import tiktoken

def load_config() -> dict:
    """加载配置文件"""
    config_path = os.path.join(os.path.dirname(__file__), 'config.yaml')
    if not os.path.exists(config_path):
        raise FileNotFoundError(f"配置文件不存在: {config_path}")
    
    with open(config_path, 'r', encoding='utf-8') as f:
        return yaml.safe_load(f)

# 加载配置
CONFIG = load_config()

class LLMClient:
    """LLM API客户端"""
    def __init__(self, config: dict):
        self.api_url = config['llm']['api_url']
        self.model_name = config['llm']['model_name']
        self.api_key = config['llm']['api_key']
        self.timeout = config['llm']['timeout']
        self.max_retries = config['llm']['max_retries']
        self.temperature = config['llm']['temperature']
        
        # 设置请求头
        self.headers = {
            "Content-Type": "application/json"
        }
        if self.api_key:
            self.headers["Authorization"] = f"Bearer {self.api_key}"
    
    def generate_response(self, messages: List[Dict[str, str]]) -> Tuple[Optional[str], Optional[str]]:
        """生成响应
        
        Args:
            messages: 消息列表
            
        Returns:
            Tuple[Optional[str], Optional[str]]: (响应内容, 错误信息)
        """
        payload = {
            "model": self.model_name,
            "messages": messages,
            "temperature": self.temperature
        }
        
        for attempt in range(self.max_retries):
            try:
                # 添加重试延迟
                if attempt > 0:
                    delay = min(2 ** attempt, 30)  # 指数退避，最大30秒
                    print(f"第 {attempt + 1} 次重试，等待 {delay} 秒...")
                    time.sleep(delay)
                
                response = requests.post(
                    self.api_url,
                    headers=self.headers,
                    json=payload,
                    timeout=self.timeout
                )
                
                # 处理不同的HTTP状态码
                if response.status_code == 200:
                    response_data = response.json()
                    return response_data['choices'][0]['message']['content'], None
                elif response.status_code == 502:
                    error_msg = f"服务器暂时不可用 (502)，正在进行第 {attempt + 1} 次重试"
                    print(error_msg)
                    if attempt < self.max_retries - 1:
                        continue
                    return None, error_msg
                elif response.status_code == 429:
                    error_msg = f"请求频率限制 (429)，正在进行第 {attempt + 1} 次重试"
                    print(error_msg)
                    if attempt < self.max_retries - 1:
                        continue
                    return None, error_msg
                elif response.status_code == 503:
                    error_msg = f"服务暂时不可用 (503)，正在进行第 {attempt + 1} 次重试"
                    print(error_msg)
                    if attempt < self.max_retries - 1:
                        continue
                    return None, error_msg
                else:
                    error_msg = f"API请求失败: HTTP {response.status_code} - {response.text}"
                    if attempt < self.max_retries - 1:
                        continue
                    return None, error_msg
                    
            except requests.Timeout:
                error_msg = f"请求超时，正在进行第 {attempt + 1} 次重试"
                print(error_msg)
                if attempt < self.max_retries - 1:
                    continue
                return None, error_msg
            except requests.ConnectionError:
                error_msg = f"连接错误，正在进行第 {attempt + 1} 次重试"
                print(error_msg)
                if attempt < self.max_retries - 1:
                    continue
                return None, error_msg
            except Exception as e:
                error_msg = f"请求出错: {str(e)}，正在进行第 {attempt + 1} 次重试"
                print(error_msg)
                if attempt < self.max_retries - 1:
                    continue
                return None, error_msg
        
        return None, f"达到最大重试次数 ({self.max_retries})"

def get_input_files() -> List[str]:
    """获取输入目录下的所有支持的文件（包括子文件夹）"""
    input_dir = CONFIG['paths']['input_dir']
    supported_extensions = CONFIG['processing']['supported_extensions']
    
    if not os.path.exists(input_dir):
        os.makedirs(input_dir)
        print(f"创建输入目录: {input_dir}")
        return []
    
    valid_files = []
    # 使用 os.walk 递归遍历所有子文件夹
    for root, _, files in os.walk(input_dir):
        for file_name in files:
            file_path = os.path.join(root, file_name)
            ext = Path(file_path).suffix.lower()
            if ext in supported_extensions:
                valid_files.append(file_path)
    
    if not valid_files:
        print(f"在 {input_dir} 目录及其子目录下没有找到支持的文件")
        print(f"支持的文件类型: {', '.join(supported_extensions)}")
        return []
    
    return valid_files

def get_output_path() -> str:
    """获取输出文件路径"""
    output_dir = CONFIG['paths']['output_dir']
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    timestamp = pd.Timestamp.now().strftime("%Y%m%d_%H%M%S")
    template = CONFIG['output']['csv_filename_template']
    return os.path.join(output_dir, template.format(timestamp=timestamp))

def extract_title(file_path: str) -> str:
    """从文件中提取标题"""
    file_ext = Path(file_path).suffix.lower()
    
    if file_ext == '.md':
        # 从Markdown文件中提取标题
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
            # 查找第一个标题
            match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
            if match:
                return match.group(1).strip()
    elif file_ext in ['.docx', '.doc']:
        try:
            # 从Word文件中提取标题
            doc = Document(file_path)
            if doc.paragraphs:
                return doc.paragraphs[0].text.strip()
        except ImportError:
            print("警告: 未安装python-docx包，无法读取Word文件")
            print("请运行: pip install python-docx")
        except Exception as e:
            print(f"读取Word文件时出错: {str(e)}")
    elif file_ext == '.pdf':
        try:
            # 从PDF文件中提取标题（使用第一页的第一行）
            reader = PdfReader(file_path)
            if reader.pages:
                first_page = reader.pages[0]
                text = first_page.extract_text()
                first_line = text.split('\n')[0].strip()
                return first_line
        except ImportError:
            print("警告: 未安装PyPDF2包，无法读取PDF文件")
            print("请运行: pip install PyPDF2")
        except Exception as e:
            print(f"读取PDF文件时出错: {str(e)}")
    elif file_ext in ['.xlsx', '.xls']:
        try:
            # 从Excel文件中提取标题（使用第一个单元格）
            if file_ext == '.xlsx':
                df = pd.read_excel(file_path)
                if not df.empty:
                    return str(df.columns[0])
            else:
                workbook = xlrd.open_workbook(file_path)
                sheet = workbook.sheet_by_index(0)
                if sheet.nrows > 0 and sheet.ncols > 0:
                    return str(sheet.cell_value(0, 0))
        except ImportError:
            print("警告: 未安装必要的包，无法读取Excel文件")
            print("请运行: pip install pandas openpyxl xlrd")
        except Exception as e:
            print(f"读取Excel文件时出错: {str(e)}")
    
    # 如果没有找到标题，使用文件名
    return Path(file_path).stem

def read_document(file_path: str) -> str:
    """读取文档内容"""
    file_ext = Path(file_path).suffix.lower()
    
    if file_ext == '.md':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    elif file_ext in ['.docx', '.doc']:
        try:
            doc = Document(file_path)
            return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        except ImportError:
            raise ImportError("未安装python-docx包，无法读取Word文件。请运行: pip install python-docx")
        except Exception as e:
            raise Exception(f"读取Word文件时出错: {str(e)}")
    elif file_ext == '.pdf':
        try:
            reader = PdfReader(file_path)
            text = []
            for page in reader.pages:
                text.append(page.extract_text())
            return '\n'.join(text)
        except ImportError:
            raise ImportError("未安装PyPDF2包，无法读取PDF文件。请运行: pip install PyPDF2")
        except Exception as e:
            raise Exception(f"读取PDF文件时出错: {str(e)}")
    elif file_ext in ['.xlsx', '.xls']:
        try:
            if file_ext == '.xlsx':
                # 使用pandas读取xlsx文件
                df = pd.read_excel(file_path)
            else:
                # 使用xlrd读取xls文件
                workbook = xlrd.open_workbook(file_path)
                sheet = workbook.sheet_by_index(0)
                data = []
                for row in range(sheet.nrows):
                    row_data = []
                    for col in range(sheet.ncols):
                        cell_value = sheet.cell_value(row, col)
                        row_data.append(str(cell_value))
                    data.append(row_data)
                df = pd.DataFrame(data)
            
            # 将DataFrame转换为文本
            text_parts = []
            # 添加列名
            text_parts.append(' | '.join(df.columns))
            # 添加数据行
            for _, row in df.iterrows():
                text_parts.append(' | '.join(str(val) for val in row))
            return '\n'.join(text_parts)
        except ImportError as e:
            raise ImportError(f"未安装必要的包，无法读取Excel文件。请运行: pip install pandas openpyxl xlrd")
        except Exception as e:
            raise Exception(f"读取Excel文件时出错: {str(e)}")
    else:
        raise ValueError(f"不支持的文件格式: {file_ext}")

def clean_for_excel(text):
    """清理Excel中不允许的字符
    
    Excel中不允许的字符包括: 
    - ASCII值小于32的控制字符（除了\t, \n, \r）
    - 某些特殊Unicode字符
    """
    if not isinstance(text, str):
        return text
        
    # 保留制表符、换行符和回车符
    allowed_control_chars = ['\t', '\n', '\r']
    
    # 清理控制字符和特殊字符
    cleaned_text = ''
    for char in text:
        # 如果是允许的控制字符，保留
        if char in allowed_control_chars:
            cleaned_text += char
            continue
            
        # 获取字符的Unicode类别
        category = unicodedata.category(char)
        
        # 排除控制字符 (Cc类别)但保留允许的控制字符
        if category == 'Cc':
            cleaned_text += ' '  # 用空格替换控制字符
        else:
            # 对于其他字符，尝试保留
            try:
                # 这里尝试处理可能导致Excel问题的字符
                # 如果是特殊字符但不是控制字符，尝试保留
                cleaned_text += char
            except:
                # 如果还有问题，就替换为空格
                cleaned_text += ' '
    
    return cleaned_text

def count_tokens(text: str, model: str = "gpt-3.5-turbo") -> int:
    """计算文本的token数量"""
    try:
        encoding = tiktoken.encoding_for_model(model)
        return len(encoding.encode(text))
    except:
        # 如果无法使用tiktoken，使用简单的估算方法
        # 假设每个token平均是4个字符
        return len(text) // 4

def split_text(text: str, max_tokens: int = 30000) -> List[str]:
    """将长文本分割成不超过最大token数的片段
    
    Args:
        text: 要分割的文本
        max_tokens: 每个片段的最大token数
        
    Returns:
        List[str]: 分割后的文本片段列表
    """
    if count_tokens(text) <= max_tokens:
        return [text]
    
    # 按段落分割文本
    paragraphs = text.split('\n\n')
    chunks = []
    current_chunk = []
    current_tokens = 0
    
    for para in paragraphs:
        para_tokens = count_tokens(para)
        
        # 如果单个段落超过最大token数，需要进一步分割
        if para_tokens > max_tokens:
            # 如果当前chunk不为空，先保存
            if current_chunk:
                chunks.append('\n\n'.join(current_chunk))
                current_chunk = []
                current_tokens = 0
            
            # 按句子分割长段落
            sentences = re.split(r'([.!?。！？])', para)
            temp_sentence = ''
            
            for i in range(0, len(sentences), 2):
                if i + 1 < len(sentences):
                    sentence = sentences[i] + sentences[i + 1]
                else:
                    sentence = sentences[i]
                
                sentence_tokens = count_tokens(sentence)
                
                if count_tokens(temp_sentence + sentence) <= max_tokens:
                    temp_sentence += sentence
                else:
                    if temp_sentence:
                        chunks.append(temp_sentence)
                    temp_sentence = sentence
            
            if temp_sentence:
                chunks.append(temp_sentence)
            continue
        
        # 检查添加当前段落是否会超过限制
        if current_tokens + para_tokens <= max_tokens:
            current_chunk.append(para)
            current_tokens += para_tokens
        else:
            # 保存当前chunk并开始新的chunk
            if current_chunk:
                chunks.append('\n\n'.join(current_chunk))
            current_chunk = [para]
            current_tokens = para_tokens
    
    # 添加最后一个chunk
    if current_chunk:
        chunks.append('\n\n'.join(current_chunk))
    
    return chunks

def generate_qa_pairs(document_content: str, llm_client: LLMClient) -> Tuple[List[Dict[str, str]], Optional[str]]:
    """使用LLM API生成问答对"""
    
    # 系统提示词
    system_prompt = """你是一个专业的文档分析助手。你的任务是：
    1. 仔细阅读提供的文档内容
    2. 提出10个与文档内容相关的重要问题
    3. 对每个问题，从文档中提取相关信息作为答案
    4. 确保问题和答案都是清晰、准确且相关的
    5. 以JSON格式返回结果，格式为：[{"question": "问题1", "answer": "答案1"}, ...]
    6. 只返回JSON格式的数据，不要包含任何其他内容"""
    
    # 用户提示词
    user_prompt = f"请分析以下文档内容并生成问答对：\n\n{document_content}"
    
    # 准备消息
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt}
    ]
    
    # 调用LLM API
    content, error = llm_client.generate_response(messages)
    if error:
        return [], error
    
    # 清理和预处理内容
    content = content.strip()
    
    # 尝试多种方式提取JSON
    qa_pairs = None
    
    # 1. 尝试直接解析整个内容
    try:
        qa_pairs = json.loads(content)
    except json.JSONDecodeError:
        pass
    
    # 2. 尝试提取JSON部分
    if qa_pairs is None:
        json_patterns = [
            r'```json\s*([\s\S]*?)\s*```',  # 匹配 ```json ... ```
            r'```\s*([\s\S]*?)\s*```',      # 匹配 ``` ... ```
            r'\[\s*\{[\s\S]*\}\s*\]'       # 匹配 [...]
        ]
        
        for pattern in json_patterns:
            match = re.search(pattern, content)
            if match:
                try:
                    json_str = match.group(1) if len(match.groups()) > 0 else match.group(0)
                    # 清理JSON字符串
                    json_str = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', json_str)  # 移除控制字符
                    json_str = re.sub(r'\\[^"\\\/bfnrtu]', '', json_str)      # 移除无效的转义序列
                    qa_pairs = json.loads(json_str)
                    break
                except json.JSONDecodeError:
                    continue
    
    # 3. 如果仍然失败，尝试手动构建JSON
    if qa_pairs is None:
        # 查找所有问题和答案对
        qa_pattern = r'问题[：:]\s*(.*?)\s*答案[：:]\s*(.*?)(?=问题[：:]|$)'
        matches = re.finditer(qa_pattern, content, re.DOTALL)
        qa_pairs = []
        for match in matches:
            question = match.group(1).strip()
            answer = match.group(2).strip()
            if question and answer:
                qa_pairs.append({
                    "question": question,
                    "answer": answer
                })
    
    # 验证结果
    if not qa_pairs or not isinstance(qa_pairs, list):
        return [], "无法解析问答对"
    
    # 验证每个问答对的结构
    valid_qa_pairs = []
    for qa in qa_pairs:
        if isinstance(qa, dict) and 'question' in qa and 'answer' in qa:
            # 清理问题和答案中的特殊字符
            qa['question'] = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', qa['question'])
            qa['answer'] = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', qa['answer'])
            valid_qa_pairs.append(qa)
    
    if not valid_qa_pairs:
        return [], "没有找到有效的问答对"
    
    return valid_qa_pairs, None

def save_failed_tasks(failed_files: List[Tuple[str, str, str]], output_dir: str):
    """保存失败任务清单到Markdown文件"""
    if not failed_files:
        return
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    failed_tasks_file = os.path.join(output_dir, f"failed_tasks_{timestamp}.md")
    
    with open(failed_tasks_file, 'w', encoding='utf-8') as f:
        f.write("# 失败任务清单\n\n")
        f.write(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
        f.write("## 失败文件列表\n\n")
        
        for file_path, error, model_response in failed_files:
            f.write(f"### {os.path.basename(file_path)}\n\n")
            f.write(f"- 文件路径: `{file_path}`\n")
            f.write(f"- 错误信息: {error}\n")
            if model_response:
                f.write("\n#### 模型返回内容\n\n")
                f.write("```\n")
                f.write(model_response)
                f.write("\n```\n\n")
        
        f.write("## 失败原因统计\n\n")
        # 统计错误类型
        error_types = {}
        for _, error, _ in failed_files:
            error_type = error.split(':')[0] if ':' in error else error
            error_types[error_type] = error_types.get(error_type, 0) + 1
        
        for error_type, count in error_types.items():
            f.write(f"- {error_type}: {count}个文件\n")
    
    print(f"\n失败任务清单已保存到: {failed_tasks_file}")

def process_files(input_files: List[str], output_file: str):
    """处理多个文件的主函数"""
    all_qa_pairs = []
    failed_files = []
    
    # 创建LLM客户端
    llm_client = LLMClient(CONFIG)
    
    # 使用tqdm创建进度条
    for file_path in tqdm(input_files, desc="处理文件"):
        try:
            # 读取文档内容
            try:
                document_content = read_document(file_path)
            except ImportError as e:
                print(f"\n处理文件 {file_path} 时出错: {str(e)}")
                failed_files.append((file_path, str(e), None))
                continue
            except Exception as e:
                print(f"\n处理文件 {file_path} 时出错: {str(e)}")
                failed_files.append((file_path, str(e), None))
                continue
            
            # 分割长文本
            text_chunks = split_text(document_content)
            file_qa_pairs = []
            
            # 处理每个文本块
            for i, chunk in enumerate(text_chunks):
                # 生成问答对
                qa_pairs, error = generate_qa_pairs(chunk, llm_client)
                
                if qa_pairs:
                    # 为每个问答对添加文件名和块序号
                    for qa in qa_pairs:
                        qa['filename'] = os.path.basename(file_path)
                        qa['chunk_index'] = i + 1
                    file_qa_pairs.extend(qa_pairs)
                else:
                    print(f"\n处理文件 {file_path} 的第 {i+1} 个文本块时出错: {error}")
            
            if file_qa_pairs:
                all_qa_pairs.extend(file_qa_pairs)
            else:
                failed_files.append((file_path, "未能生成任何问答对", document_content))
            
        except Exception as e:
            print(f"\n处理文件 {file_path} 时发生错误: {str(e)}")
            failed_files.append((file_path, str(e), None))
            continue
    
    if not all_qa_pairs:
        print("\n未能生成任何问答对")
        if failed_files:
            print("\n失败的文件列表：")
            for file_path, error, _ in failed_files:
                print(f"- {file_path}: {error}")
        return
    
    # 保存为CSV
    df = pd.DataFrame(all_qa_pairs)
    df = df[['filename', 'chunk_index', 'question', 'answer']]  # 重新排序列
    df.to_csv(output_file, index=False, encoding='utf-8', escapechar='\\', quoting=csv.QUOTE_ALL)
    
    # 清理Excel中不允许的字符
    for col in df.columns:
        df[col] = df[col].apply(clean_for_excel)
    
    # 保存为Excel
    try:
        timestamp = pd.Timestamp.now().strftime("%Y%m%d_%H%M%S")
        excel_template = CONFIG['output']['excel_filename_template']
        excel_file = os.path.join(CONFIG['paths']['output_dir'], 
                                excel_template.format(timestamp=timestamp))
        df.to_excel(excel_file, index=False)
        print(f"Excel文件已保存到: {excel_file}")
    except Exception as e:
        print(f"保存Excel文件时出错: {str(e)}")
        print("仅保存为CSV格式")
    
    print(f"\n处理完成！")
    print(f"CSV文件已保存到: {output_file}")
    print(f"共处理 {len(input_files)} 个文件，成功处理 {len(input_files) - len(failed_files)} 个文件，生成 {len(all_qa_pairs)} 个问答对")
    
    # 保存失败任务清单
    if failed_files:
        save_failed_tasks(failed_files, CONFIG['paths']['output_dir'])

if __name__ == "__main__":
    print(f"正在扫描目录: {CONFIG['paths']['input_dir']}")
    input_files = get_input_files()
    
    if not input_files:
        print("没有找到要处理的文件，程序退出")
        exit(1)
    
    print(f"找到 {len(input_files)} 个文件")
    output_file = get_output_path()
    
    # 处理文件
    process_files(input_files, output_file) 